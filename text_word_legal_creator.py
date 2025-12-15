#!/usr/bin/env python3
"""
OCRデータを基にした役員変更登記申請書（テキスト・Word形式）作成システム
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class TextWordLegalCreator:
    def __init__(self):
        self.output_dir = "/home/fujinosuke/ocr_inbox"
        
    def create_text_document(self):
        """テキスト形式の申請書を作成"""
        
        today = datetime.now()
        
        text_content = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                    役員変更登記申請書
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


１．商　　　号　　有限会社越後屋商店


１．本　　　店　　北海道札幌市南区川沿十三条二丁目1番51号


１．登記の事由　　取締役の変更


１．登記すべき事項
　　　　　　　　　平成20年10月20日
　　　　　　　　　　　北海道札幌市中央区宮の森一条十五丁目5番12―305号
　　　　　　　　　　　　　取締役荒井　尚辞任


１．登録免許税　　金１０，０００円


１．添付書類
　　　　　　　　　辞任届　　　　　　　　　　　１通




　　　　　　上記のとおり登記の申請をします。


　　　　　　　　　　　　　　　　　　　　　　　平成20年10月23日


　　　　　　　　　　　　北海道札幌市南区川沿十三条二丁目1番51号
　　　　　　　　　　　　　　申請人　有限会社越後屋商店

　　　　　　　　　　　　北海道札幌市南区川沿十三条二丁目1番51号
　　　　　　　　　　　　　　代表取締役　佐藤　明美　　　　　[印]

　　　　　　　　　　　　連絡先の電話番号　（〇一一）五七三―〇七四〇




　　　　　　　　　　　　　　札幌法務局　御中


━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                    収入印紙貼付台紙
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


　　　　　　　　　　　　　　┌─────────────┐
　　　　　　　　　　　　　　│                     │
　　　　　　　　　　　　　　│     収　入　印　紙     │
　　　　　　　　　　　　　　│                     │
　　　　　　　　　　　　　　│     １０，０００円     │
　　　　　　　　　　　　　　│                     │
　　　　　　　　　　　　　　│   ※貼付・割印要     │
　　　　　　　　　　　　　　│                     │
　　　　　　　　　　　　　　└─────────────┘


━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                         辞　任　届
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


　私は，このたび一身上の都合により，貴社の取締役を辞任いたしたく，
お届けいたします。


　　　　　　　　　　　　　　　　　　　　　　　平成20年10月20日


　　　　　　　　　　北海道札幌市中央区宮の森一条十五丁目5番12―305号
　　　　　　　　　　　　　　　　　　　　　　　　　荒井　尚　　　[印]




　　　　　　　　　　　　有限会社越後屋商店　御中


━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                    受付番号票貼付欄
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


　　　　　　　　　　　　　　┌─────────────┐
　　　　　　　　　　　　　　│                     │
　　　　　　　　　　　　　　│  受付番号票貼付欄    │
　　　　　　　　　　　　　　│                     │
　　　　　　　　　　　　　　│ ※法務局で発行される │
　　　　　　　　　　　　　　│ 受付番号票を貼付    │
　　　　　　　　　　　　　　│                     │
　　　　　　　　　　　　　　└─────────────┘


【法的根拠】
・商業登記法第20条（登記の申請）
・商業登記法第21条（申請書の記載事項）
・商業登記法第24条（添付書面）
・会社法第911条第3項第3号（役員の変更登記）

【提出書類一覧】
１．役員変更登記申請書　　　　　　　１通
２．辞任届（荒井尚）　　　　　　　　１通
３．収入印紙貼付台紙　　　　　　　　１通
４．受付番号票貼付欄　　　　　　　　１通

作成日: {today.strftime('%Y年%m月%d日')}
元データ: 11-1-07法務局役員変更.rtf（OCR読み取り）
作成システム: Claude Code Assistant
"""
        
        # テキストファイルを保存
        timestamp = today.strftime('%Y%m%d_%H%M%S')
        filename = f"legal_document_text_{timestamp}.txt"
        file_path = os.path.join(self.output_dir, filename)
        
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            print(f"✅ テキスト形式申請書を作成しました: {filename}")
            print(f"📁 保存場所: {file_path}")
            
            return file_path
            
        except Exception as e:
            print(f"❌ テキストファイル保存に失敗: {e}")
            return None
    
    def create_word_document(self):
        """Word形式の申請書を作成"""
        
        today = datetime.now()
        
        # 新しいWordドキュメントを作成
        doc = Document()
        
        # ページ設定（A4）
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # ===== 1ページ目: 役員変更登記申請書 =====
        
        # タイトル
        title = doc.add_heading('', level=0)
        title_run = title.runs[0] if title.runs else title.add_run()
        title_run.text = '役員変更登記申請書'
        title_run.font.size = Pt(20)
        title_run.font.name = 'MS Gothic'
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('')  # 空行
        
        # 申請書本文
        items = [
            ('１．商　　　号', '有限会社越後屋商店'),
            ('１．本　　　店', '北海道札幌市南区川沿十三条二丁目1番51号'),
            ('１．登記の事由', '取締役の変更'),
            ('１．登録免許税', '金１０，０００円')
        ]
        
        for label, content in items:
            p = doc.add_paragraph()
            p.add_run(label).font.name = 'MS Gothic'
            p.add_run('\t').font.name = 'MS Gothic'
            p.add_run(content).font.name = 'MS Gothic'
            doc.add_paragraph('')  # 空行
        
        # 登記すべき事項（特別レイアウト）
        p = doc.add_paragraph()
        p.add_run('１．登記すべき事項').font.name = 'MS Gothic'
        
        p2 = doc.add_paragraph()
        p2.add_run('\t\t\t\t平成20年10月20日').font.name = 'MS Gothic'
        
        p3 = doc.add_paragraph()
        p3.add_run('\t\t\t\t\t北海道札幌市中央区宮の森一条十五丁目5番12―305号').font.name = 'MS Gothic'
        
        p4 = doc.add_paragraph()
        p4.add_run('\t\t\t\t\t\t取締役荒井　尚辞任').font.name = 'MS Gothic'
        
        doc.add_paragraph('')
        
        # 添付書類
        p = doc.add_paragraph()
        p.add_run('１．添付書類').font.name = 'MS Gothic'
        
        p2 = doc.add_paragraph()
        p2.add_run('\t\t\t\t辞任届\t\t\t\t\t１通').font.name = 'MS Gothic'
        
        # 申請文
        for _ in range(3):
            doc.add_paragraph('')
        
        p = doc.add_paragraph('上記のとおり登記の申請をします。')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = 'MS Gothic'
        
        doc.add_paragraph('')
        doc.add_paragraph('')
        
        # 日付
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run('平成20年10月23日').font.name = 'MS Gothic'
        
        doc.add_paragraph('')
        doc.add_paragraph('')
        
        # 申請者情報
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run('北海道札幌市南区川沿十三条二丁目1番51号').font.name = 'MS Gothic'
        
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.add_run('申請人　有限会社越後屋商店').font.name = 'MS Gothic'
        
        doc.add_paragraph('')
        
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.add_run('北海道札幌市南区川沿十三条二丁目1番51号').font.name = 'MS Gothic'
        
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p4.add_run('代表取締役　佐藤　明美　　　　　[印]').font.name = 'MS Gothic'
        
        doc.add_paragraph('')
        
        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p5.add_run('連絡先の電話番号　（〇一一）五七三―〇七四〇').font.name = 'MS Gothic'
        
        for _ in range(4):
            doc.add_paragraph('')
        
        # 宛先
        p = doc.add_paragraph('札幌法務局　御中')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = 'MS Gothic'
        p.runs[0].bold = True
        
        # ===== 2ページ目: 収入印紙貼付台紙 =====
        doc.add_page_break()
        
        title2 = doc.add_heading('', level=0)
        title2_run = title2.add_run('収入印紙貼付台紙')
        title2_run.font.size = Pt(20)
        title2_run.font.name = 'MS Gothic'
        title2_run.bold = True
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for _ in range(5):
            doc.add_paragraph('')
        
        # 印紙枠（テーブルで表現）
        table = doc.add_table(rows=3, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        cell1 = table.rows[0].cells[0]
        cell1.text = '収　入　印　紙'
        cell1_p = cell1.paragraphs[0]
        cell1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1_p.runs[0].font.name = 'MS Gothic'
        cell1_p.runs[0].bold = True
        
        cell2 = table.rows[1].cells[0]
        cell2.text = '１０，０００円'
        cell2_p = cell2.paragraphs[0]
        cell2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2_p.runs[0].font.name = 'MS Gothic'
        cell2_p.runs[0].bold = True
        
        cell3 = table.rows[2].cells[0]
        cell3.text = '※貼付・割印要'
        cell3_p = cell3.paragraphs[0]
        cell3_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell3_p.runs[0].font.name = 'MS Gothic'
        cell3_p.runs[0].font.size = Pt(10)
        
        # ===== 3ページ目: 辞任届 =====
        doc.add_page_break()
        
        title3 = doc.add_heading('', level=0)
        title3_run = title3.add_run('辞　任　届')
        title3_run.font.size = Pt(20)
        title3_run.font.name = 'MS Gothic'
        title3_run.bold = True
        title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for _ in range(3):
            doc.add_paragraph('')
        
        # 辞任届本文
        p = doc.add_paragraph('　私は，このたび一身上の都合により，貴社の取締役を辞任いたしたく，お届けいたします。')
        p.runs[0].font.name = 'MS Gothic'
        
        for _ in range(5):
            doc.add_paragraph('')
        
        # 日付
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run('平成20年10月20日').font.name = 'MS Gothic'
        
        for _ in range(3):
            doc.add_paragraph('')
        
        # 辞任者情報
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run('北海道札幌市中央区宮の森一条十五丁目5番12―305号').font.name = 'MS Gothic'
        
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.add_run('荒井　尚　　　[印]').font.name = 'MS Gothic'
        
        for _ in range(5):
            doc.add_paragraph('')
        
        # 宛先
        p = doc.add_paragraph('有限会社越後屋商店　御中')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.runs[0].font.name = 'MS Gothic'
        
        # ===== 4ページ目: 受付番号票貼付欄・法的根拠 =====
        doc.add_page_break()
        
        title4 = doc.add_heading('', level=0)
        title4_run = title4.add_run('受付番号票貼付欄')
        title4_run.font.size = Pt(18)
        title4_run.font.name = 'MS Gothic'
        title4_run.bold = True
        title4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for _ in range(3):
            doc.add_paragraph('')
        
        # 受付番号票枠
        table2 = doc.add_table(rows=2, cols=1)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        cell1 = table2.rows[0].cells[0]
        cell1.text = '受付番号票貼付欄'
        cell1_p = cell1.paragraphs[0]
        cell1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1_p.runs[0].font.name = 'MS Gothic'
        cell1_p.runs[0].bold = True
        
        cell2 = table2.rows[1].cells[0]
        cell2.text = '※法務局で発行される受付番号票を貼付'
        cell2_p = cell2.paragraphs[0]
        cell2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2_p.runs[0].font.name = 'MS Gothic'
        cell2_p.runs[0].font.size = Pt(10)
        
        for _ in range(3):
            doc.add_paragraph('')
        
        # 法的根拠
        p = doc.add_paragraph()
        p.add_run('【法的根拠】').font.name = 'MS Gothic'
        p.runs[0].bold = True
        
        legal_items = [
            '・商業登記法第20条（登記の申請）',
            '・商業登記法第21条（申請書の記載事項）',
            '・商業登記法第24条（添付書面）',
            '・会社法第911条第3項第3号（役員の変更登記）'
        ]
        
        for item in legal_items:
            p = doc.add_paragraph(item)
            p.runs[0].font.name = 'MS Gothic'
            p.runs[0].font.size = Pt(10)
        
        doc.add_paragraph('')
        
        # 提出書類一覧
        p = doc.add_paragraph()
        p.add_run('【提出書類一覧】').font.name = 'MS Gothic'
        p.runs[0].bold = True
        
        doc_items = [
            '１．役員変更登記申請書　　　　　　　１通',
            '２．辞任届（荒井尚）　　　　　　　　１通',
            '３．収入印紙貼付台紙　　　　　　　　１通',
            '４．受付番号票貼付欄　　　　　　　　１通'
        ]
        
        for item in doc_items:
            p = doc.add_paragraph(item)
            p.runs[0].font.name = 'MS Gothic'
            p.runs[0].font.size = Pt(10)
        
        doc.add_paragraph('')
        
        # 作成情報
        p = doc.add_paragraph(f'作成日: {today.strftime("%Y年%m月%d日")}')
        p.runs[0].font.name = 'MS Gothic'
        p.runs[0].font.size = Pt(8)
        
        p = doc.add_paragraph('元データ: 11-1-07法務局役員変更.rtf（OCR読み取り）')
        p.runs[0].font.name = 'MS Gothic'
        p.runs[0].font.size = Pt(8)
        
        p = doc.add_paragraph('作成システム: Claude Code Assistant')
        p.runs[0].font.name = 'MS Gothic'
        p.runs[0].font.size = Pt(8)
        
        # Wordファイルを保存
        timestamp = today.strftime('%Y%m%d_%H%M%S')
        filename = f"legal_document_word_{timestamp}.docx"
        file_path = os.path.join(self.output_dir, filename)
        
        try:
            doc.save(file_path)
            
            print(f"✅ Word形式申請書を作成しました: {filename}")
            print(f"📁 保存場所: {file_path}")
            
            return file_path
            
        except Exception as e:
            print(f"❌ Wordファイル保存に失敗: {e}")
            return None

def main():
    """メイン処理"""
    print("📄 OCRデータ基準テキスト・Word形式申請書作成システム")
    print("="*70)
    print("📋 OCRデータ: 有限会社越後屋商店（11-1-07法務局役員変更.rtf）")
    print("="*70)
    
    creator = TextWordLegalCreator()
    
    # テキスト形式作成
    print("\n📝 テキスト形式申請書を作成中...")
    text_file = creator.create_text_document()
    
    # Word形式作成
    print("\n📄 Word形式申請書を作成中...")
    word_file = creator.create_word_document()
    
    if text_file and word_file:
        print("\n" + "="*70)
        print("✅ テキスト・Word形式申請書の作成が完了しました")
        print("="*70)
        print(f"📝 テキストファイル: {os.path.basename(text_file)}")
        print(f"📄 Wordファイル: {os.path.basename(word_file)}")
        
        print("\n📋 実際のOCRデータを使用した内容:")
        print("• 🏢 会社名: 有限会社越後屋商店")
        print("• 📍 本店: 北海道札幌市南区川沿十三条二丁目1番51号")
        print("• 👤 退任者: 荒井尚（取締役）")
        print("• 📅 退任日: 平成20年10月20日")
        print("• 💰 登録免許税: 10,000円")
        print("• 📞 連絡先: (011)573-0740")
        print("• 👩‍💼 代表取締役: 佐藤明美")
        
        print("\n📄 作成された書類（両形式共通）:")
        print("• 📋 役員変更登記申請書（1ページ目）")
        print("• 💴 収入印紙貼付台紙（2ページ目）")
        print("• 📝 辞任届（3ページ目）")
        print("• 🎫 受付番号票貼付欄（4ページ目）")
        
        print("\n📄 ファイル形式の特徴:")
        print("📝 テキスト形式:")
        print("  • シンプルなテキスト表示")
        print("  • どの環境でも開ける")
        print("  • 軽量ファイル")
        print("  • ASCII文字による罫線")
        
        print("\n📄 Word形式:")
        print("  • 正式なレイアウト")
        print("  • 印刷最適化")
        print("  • 編集可能")
        print("  • MS Gothic標準フォント")
        print("  • 4ページ構成")
        
        print(f"\n🔧 使用方法:")
        print(f"  📝 テキスト: メモ帳等で {text_file} を開く")
        print(f"  📄 Word: Microsoft Word等で {word_file} を開く")
        
        return True
    else:
        print("❌ ドキュメント作成に失敗しました")
        return False

if __name__ == "__main__":
    main()